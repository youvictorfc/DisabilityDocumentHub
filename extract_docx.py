import docx
import json

def extract_docx_content(file_path):
    doc = docx.Document(file_path)
    
    # Extract text from paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    
    # Extract text from tables
    tables = []
    for table in doc.tables:
        table_data = []
        for i, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                # Get text from each cell
                cell_text = ""
                for paragraph in cell.paragraphs:
                    cell_text += paragraph.text + "\n"
                row_data.append(cell_text.strip())
            table_data.append(row_data)
        tables.append(table_data)
    
    return {
        "paragraphs": paragraphs,
        "tables": tables
    }

if __name__ == "__main__":
    file_path = "attached_assets/Home Safety Checklist.docx"
    content = extract_docx_content(file_path)
    
    # Print paragraphs
    print("=== PARAGRAPHS ===")
    for i, para in enumerate(content["paragraphs"]):
        print(f"{i+1}. {para}")
    
    # Print tables
    print("\n=== TABLES ===")
    for i, table in enumerate(content["tables"]):
        print(f"Table {i+1}:")
        for row in table:
            print(row)
        print()