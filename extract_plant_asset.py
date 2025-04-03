import docx
import json

def extract_docx_content(file_path):
    """Extract content from a DOCX file."""
    doc = docx.Document(file_path)
    
    # Extract paragraphs
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    
    # Extract tables
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)
        tables.append(table_data)
    
    return {"paragraphs": paragraphs, "tables": tables}

if __name__ == "__main__":
    file_path = "attached_assets/New Plant-Asset Hazard Checklist.docx"
    content = extract_docx_content(file_path)
    print(json.dumps(content, indent=2))