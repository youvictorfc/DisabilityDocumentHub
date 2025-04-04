"""
Debug script to test the Root Cause Analysis Form detection.
"""

import os
import sys
import docx
from services.form.root_cause_analysis_template import is_root_cause_analysis, get_root_cause_analysis_template

def main():
    # Check if the file path is provided
    if len(sys.argv) < 2:
        print("Usage: python debug_root_cause_form.py <path_to_form>")
        return
    
    file_path = sys.argv[1]
    
    # Check if the file exists
    if not os.path.exists(file_path):
        print(f"Error: File {file_path} does not exist")
        return
    
    # Basic file information
    print(f"Analyzing file: {file_path}")
    print(f"File size: {os.path.getsize(file_path)} bytes")
    file_name = os.path.basename(file_path)
    print(f"File name: {file_name}")
    
    # Keyword detection in filename
    keywords = ["root cause", "analysis", "rca"]
    detected_keywords = [kw for kw in keywords if kw in file_name.lower()]
    print(f"Keywords detected in filename: {detected_keywords}")
    
    # Check DOCX structure
    if file_path.lower().endswith('.docx'):
        try:
            doc = docx.Document(file_path)
            print(f"DOCX information:")
            print(f"Number of paragraphs: {len(doc.paragraphs)}")
            print(f"Number of tables: {len(doc.tables)}")
            
            # Print the first few paragraphs
            print("\nFirst paragraphs:")
            for i, para in enumerate(doc.paragraphs[:5]):
                if para.text.strip():
                    print(f"{i+1}: {para.text[:100]}...")
            
            # Print table information
            if doc.tables:
                print("\nTable information:")
                for i, table in enumerate(doc.tables[:3]):
                    print(f"Table {i+1}: {len(table.rows)} rows x {len(table.columns)} columns")
                    if i == 0 and len(table.rows) > 0:
                        # Print the first row of the first table
                        first_row = [cell.text for cell in table.rows[0].cells]
                        print(f"First row: {first_row}")
        except Exception as e:
            print(f"Error processing DOCX file: {str(e)}")
    
    # Test the detection function
    print("\nDetection result:")
    is_rca = is_root_cause_analysis(file_path)
    print(f"is_root_cause_analysis() result: {is_rca}")
    
    # Print the template
    print("\nTemplate structure:")
    template = get_root_cause_analysis_template()
    print(f"Number of fields in template: {len(template)}")
    print("First few fields:")
    for i, field in enumerate(template[:5]):
        print(f"{i+1}: {field['question_text']}")

if __name__ == "__main__":
    main()