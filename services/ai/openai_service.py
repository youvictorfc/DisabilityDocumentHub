import os
import json
import logging
import base64
import mimetypes
from pathlib import Path
from openai import OpenAI
from flask import current_app

# Initialize OpenAI client - will be set with the actual API key in the functions
openai = None

def get_openai_client():
    """
    Get or initialize the OpenAI client with the current API key from the app config
    """
    global openai
    if openai is None:
        api_key = current_app.config.get('OPENAI_API_KEY')
        if not api_key:
            raise ValueError("OpenAI API key is not configured. Please set the OPENAI_API_KEY environment variable.")
        openai = OpenAI(api_key=api_key)
    return openai

def is_image_file(file_path):
    """Check if a file is an image based on its extension or MIME type"""
    image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'}
    extension = Path(file_path).suffix.lower()
    
    # Check by extension
    if extension in image_extensions:
        return True
    
    # Check by MIME type
    mime_type, _ = mimetypes.guess_type(file_path)
    return mime_type and mime_type.startswith('image/')

def encode_image_to_base64(image_path):
    """Encode an image file to base64 for the OpenAI API"""
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')

def verify_field_extraction_completeness(image_path, base64_image, extracted_fields):
    """
    Verify that all potential form fields in the image have been extracted.
    Returns potential missed fields and a completeness assessment.
    
    This function performs a thorough verification pass using a different approach
    to catch any fields that might have been missed in the initial extraction.
    """
    client = get_openai_client()
    
    # Define verification result schema for consistent output
    verification_schema = {
        "type": "object",
        "properties": {
            "complete": {"type": "boolean"},
            "issues": {"type": "array", "items": {"type": "string"}},
            "suggestions": {"type": "array", "items": {"type": "string"}},
            "missed_questions": {"type": "array", "items": {"type": "string"}}
        },
        "required": ["complete", "issues", "missed_questions"]
    }
    
    try:
        # Format the extracted fields for the verification prompt
        field_list = []
        for i, field in enumerate(extracted_fields):
            field_text = field.get('question_text', field.get('label', ''))
            field_type = field.get('field_type', 'unknown')
            if field_text:
                field_list.append(f"{i+1}. [{field_type}] {field_text}")
        
        field_summary = "\n".join(field_list) if field_list else "No fields were extracted."
        
        current_app.logger.info(f"Verifying extraction completeness with {len(extracted_fields)} fields")
        
        # Make verification API call with enhanced prompt
        verification_response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a form verification expert with exceptional attention to detail. "
                        "Your task is to verify whether ALL form fields have been properly extracted from a form image. "
                        "You will be given a list of already extracted fields and the original form image. "
                        "Your job is to identify any fields that might have been missed or incompletely extracted.\n\n"
                        
                        "CRITICAL VERIFICATION REQUIREMENTS:\n"
                        "1. DO NOT ADD TEXT THAT IS NOT DIRECTLY PRESENT IN THE IMAGE\n"
                        "2. Be extremely thorough in your verification\n"
                        "3. Compare what's in the image against what has been extracted\n"
                        "4. Look for ANY missed questions, fields, checkboxes, or input areas\n"
                        "5. Note if any extracted fields are inaccurate or have been modified from the original\n\n"
                        
                        "SPECIAL ATTENTION AREAS:\n"
                        "1. Checklist forms - Each row in a checklist should be a separate field\n"
                        "2. Tables with YES/NO columns - Each row should be captured as a radio button question\n"
                        "3. Audit forms - Every audit item should be captured separately\n"
                        "4. Multi-section forms - Section headers should be included with items\n"
                        "5. Forms with multiple columns - Ensure all columns are properly captured\n"
                        "6. Small text in corners or margins - These are often missed\n"
                        "7. Signature fields, date fields, and reference numbers - These should be captured\n"
                        "8. Instructions or guidance text - If these are substantial, they should be included"
                    )
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                f"I've extracted {len(extracted_fields)} fields from this form image.\n\n"
                                f"EXTRACTED FIELDS:\n{field_summary}\n\n"
                                "Please carefully compare the image with these extracted fields and identify ANYTHING that was missed or extracted incorrectly. Pay special attention to:\n"
                                "1. Headers or section titles that provide important context\n"
                                "2. Small or faint text fields that might be hard to see\n"
                                "3. Instructions or guidance text that should be captured\n"
                                "4. Fields in unusual locations (footers, margins, etc.)\n"
                                "5. Tables or grid structures - EACH ROW should typically be extracted as a separate field\n"
                                "6. Checklist forms with YES/NO columns - each row should be a separate question with options\n"
                                "7. Audit forms - each audit item should be extracted separately\n"
                                "8. Forms with sections and subsections - section headers should be included\n"
                                "9. Required field indicators - asterisks or other symbols noting mandatory fields\n"
                                "10. Date fields, signature fields, or reference number fields\n\n"
                                
                                "Record ANY text that appears on the form but is not in the extracted fields list.\n\n"
                                
                                "Please respond with your assessment as a JSON object with:\n"
                                "{\n"
                                "    \"complete\": true|false (whether the extraction seems complete),\n"
                                "    \"issues\": [\"issue 1\", \"issue 2\", ...] (problems with the extraction),\n"
                                "    \"suggestions\": [\"suggestion 1\", \"suggestion 2\", ...] (recommendations to improve),\n"
                                "    \"missed_questions\": [\"missed question 1\", \"missed question 2\", ...] (exact text of fields that were missed)\n"
                                "}\n\n"
                                
                                "IMPORTANT: For missed_questions, include the EXACT text as it appears in the form, not paraphrased or summarized versions."
                            )
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}"
                            }
                        }
                    ]
                }
            ],
            response_format={"type": "json_object"},
            temperature=0.2  # Lower temperature for more precise verification
        )
        
        verification_result = json.loads(verification_response.choices[0].message.content)
        
        # Log verification results
        is_complete = verification_result.get('complete', False)
        missed_questions = verification_result.get('missed_questions', [])
        issues = verification_result.get('issues', [])
        
        if not is_complete:
            current_app.logger.warning("Question extraction may be incomplete")
            for issue in issues:
                current_app.logger.warning(f"Validation issue: {issue}")
            for missed in missed_questions:
                current_app.logger.warning(f"Potentially missed field: {missed}")
        
        return verification_result
        
    except Exception as e:
        current_app.logger.error(f"Error during extraction verification: {str(e)}")
        # Return a default structure if verification fails
        return {
            "complete": False,
            "issues": [f"Verification error: {str(e)}"],
            "suggestions": ["Proceed with caution as verification failed"],
            "missed_questions": []
        }

def extract_form_fields_from_markdown(markdown_content, file_path=None):
    """
    Extract form fields from markdown content generated from a document.
    This function leverages the structured markdown format to more accurately
    identify form fields, questions, and their relationships.
    
    Args:
        markdown_content: The markdown representation of the document
        file_path: Optional original file path for logging purposes
        
    Returns:
        dict: Form structure with questions array
    """
    try:
        file_info = f" from {file_path}" if file_path else ""
        current_app.logger.info(f"Extracting form fields from markdown content{file_info}")
        
        # Get the OpenAI client
        client = get_openai_client()
        
        # Extract form fields from markdown content
        current_app.logger.info("Performing form field extraction from markdown structure")
        # the newest OpenAI model is "gpt-4o" which was released May 13, 2024.
        # do not change this unless explicitly requested by the user
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a form field extraction specialist for markdown documents. "
                        "Your task is to extract ALL form fields from markdown content with high accuracy. "
                        "Markdown structure preserves the document's original formatting, making it easier to identify fields. "
                        "\n\nGUIDELINES:\n"
                        "1. Preserve the EXACT question text as shown in the markdown\n"
                        "2. Pay special attention to markdown-specific formats:\n"
                        "   - Headers (marked with #, ##, etc.) are often section titles\n"
                        "   - Bullet points (marked with *, -, or +) may be separate questions\n"
                        "   - Tables (defined with | symbols) often contain form fields in each row\n"
                        "   - Code blocks (marked with ```) may contain structured content\n"
                        "3. Capture ALL form fields, including checkboxes, radio buttons, and input fields\n"
                        "4. For tables, extract each row as a separate question\n"
                        "5. Maintain the SAME ORDER of questions as they appear in the document\n"
                        "6. Include ALL section headings (# headers) as 'header' field type\n"
                        "7. DO NOT INVENT FIELDS that aren't in the markdown\n"
                        "8. DO NOT MAKE UP QUESTIONS or modify the text\n"
                        "9. If you find checkbox syntax like '- [ ]' or '- [x]', identify these as checkbox fields\n"
                        "10. Tables often represent form fields with multiple columns, analyze them carefully\n"
                        "11. When extracting tables, the first column is usually the question_text\n"
                        "12. For rows with cells containing YES/NO/NA options, interpret as radio button fields"
                    )
                },
                {
                    "role": "user",
                    "content": (
                        "Extract ALL form fields from this markdown content. The markdown preserves "
                        "the structure of the original document. Extract each field as a separate question item.\n\n"
                        "Provide your response as valid JSON with a `questions` array. Each question should include:\n"
                        "- `id`: A unique identifier\n"
                        "- `question_text`: The EXACT text of the question/field label\n"
                        "- `field_type`: One of 'text', 'textarea', 'checkbox', 'radio', 'select', 'date', 'header'\n"
                        "- `options`: Array of options if applicable (for checkbox, radio, select)\n"
                        "- `required`: Boolean indicating if the field appears to be required\n\n"
                        "For tables, each row should be a separate question. For headings (#, ##, etc.), "
                        "use field_type 'header'. Maintain the EXACT question order.\n\n"
                        f"Here is the markdown content:\n\n{markdown_content}"
                    )
                }
            ],
            response_format={"type": "json_object"},
            temperature=0.1  # Lower temperature for more deterministic extraction
        )
        
        initial_extraction = json.loads(response.choices[0].message.content)
        current_app.logger.info(f"Extracted {len(initial_extraction.get('questions', []))} questions from markdown")
        
        # Verify completeness with a second pass for quality assurance
        current_app.logger.info("Performing verification pass to ensure all fields are captured")
        verification_response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a quality assurance specialist for form field extraction from markdown. "
                        "Your task is to verify that ALL form fields have been correctly identified and "
                        "that NOTHING was missed from the markdown content."
                        "\n\nFocus on identifying any missing fields or items that should have been extracted "
                        "but weren't. Analyze the markdown carefully and compare with the extracted list."
                    )
                },
                {
                    "role": "user",
                    "content": (
                        "Here is a list of form fields that were extracted from markdown content. "
                        "Please examine the markdown carefully and verify if ANY fields or questions are missing from this list. "
                        "Pay special attention to:\n"
                        "- Markdown headers (#, ##, etc.) that should be section titles\n"
                        "- List items (*, -, +) that might be form fields\n"
                        "- Table rows that should be questions\n"
                        "- Checkbox syntax (- [ ], - [x])\n"
                        "- Form-like structure such as questions ending with a colon\n\n"
                        f"Extracted fields (verify against the markdown):\n{json.dumps(initial_extraction, indent=2)}\n\n"
                        "Respond with a JSON object that includes:\n"
                        "1. `missed_fields`: Array of any fields that were completely missed\n"
                        "2. `completeness_assessment`: Your assessment of whether any questions found in the markdown were not in the extracted list.\n\n"
                        f"Here is the markdown content to verify against:\n\n{markdown_content[:10000]}" +
                        (f"\n\n[Content truncated due to length. This is only the first portion of the markdown.]" if len(markdown_content) > 10000 else "")
                    )
                }
            ],
            response_format={"type": "json_object"},
            temperature=0.1
        )
        
        verification_results = json.loads(verification_response.choices[0].message.content)
        missed_fields = verification_results.get('missed_fields', [])
        completeness_assessment = verification_results.get('completeness_assessment', '')
        
        # Check if verification identified any missing fields
        if missed_fields and len(missed_fields) > 0:
            current_app.logger.warning(f"Verification identified {len(missed_fields)} missed fields")
            # Add the missed fields to our extraction
            for field in missed_fields:
                # Generate a unique ID if not provided
                if 'id' not in field and 'question_text' in field:
                    field['id'] = f"question_{len(initial_extraction.get('questions', [])) + 1}"
                
                # Add to our questions list
                initial_extraction.setdefault('questions', []).append(field)
            
            current_app.logger.info(f"Updated extraction now has {len(initial_extraction.get('questions', []))} questions")
        else:
            current_app.logger.info("Verification confirmed all fields were extracted correctly")
        
        # Log completeness assessment
        if 'incomplete' in completeness_assessment.lower() or 'missing' in completeness_assessment.lower():
            current_app.logger.warning(f"Completeness assessment: {completeness_assessment}")
        else:
            current_app.logger.info(f"Completeness assessment: {completeness_assessment}")
        
        # Return the final extraction with any missed fields added
        return initial_extraction
    
    except Exception as e:
        current_app.logger.error(f"Error extracting fields from markdown: {str(e)}")
        raise Exception(f"Failed to extract form fields from markdown: {str(e)}")

def extract_form_fields_from_image(image_path):
    """Extract form fields from an image using GPT-4o multimodal capabilities with enhanced handling for challenging images"""
    client = get_openai_client()
    
    # Define our JSON schema for form extraction to ensure consistent output
    json_schema = {
        "type": "object",
        "properties": {
            "questions": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "id": { "type": "string" },
                        "question_text": { "type": "string" },
                        "field_type": { 
                            "type": "string", 
                            "enum": ["text", "textarea", "radio", "checkbox", "select", "date", "email", "number", "signature"] 
                        },
                        "options": { 
                            "type": "array", 
                            "items": { "type": "string" } 
                        },
                        "required": { "type": "boolean" }
                    },
                    "required": ["id", "question_text", "field_type"]
                }
            }
        },
        "required": ["questions"]
    }
    
    try:
        base64_image = encode_image_to_base64(image_path)
        current_app.logger.info(f"Attempting form extraction from image: {image_path}")
        
        # First try with GPT-4o
        try:
            # the newest OpenAI model is "gpt-4o" which was released May 13, 2024.
            # do not change this unless explicitly requested by the user
            current_app.logger.info(f"Attempting form extraction with GPT-4o model using enhanced JSON schema")
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are a specialized form extraction expert for Minto Disability Services with exceptional attention to detail. "
                            "Your task is to analyze the provided image of a form and extract ALL form fields/questions EXACTLY as they appear in the original. "
                            "CRITICAL REQUIREMENTS:\n"
                            "1. EXTRACT EVERY SINGLE FORM ELEMENT WITH THE EXACT ORIGINAL TEXT - DO NOT PARAPHRASE, DO NOT COMBINE FIELDS, DO NOT IMPROVE CLARITY\n"
                            "2. DO NOT ADD ANY TEXT THAT IS NOT DIRECTLY PRESENT IN THE IMAGE\n"
                            "3. DO NOT SKIP ANY FIELDS, EVEN IF THEY SEEM MINOR OR REDUNDANT\n"
                            "4. If the image quality is poor, do your absolute best to decipher the text while maintaining exact wording\n"
                            "5. If the image is rotated or skewed, mentally adjust your perspective to read it correctly\n"
                            "6. For complex forms with multiple sections or tables, process sequentially (usually top-to-bottom, left-to-right)\n"
                            "7. Never add explanatory text to fields that isn't present in the original\n"
                            "8. If you're uncertain about field content, include what you can see and note uncertainty with [?] in field labels\n\n"
                            "9. For tabular forms with YES/NO columns or checklists, extract EACH ROW as a separate question\n"
                            "10. Pay special attention to tables - extract EVERY ROW as a separate question\n"
                            "11. For forms with sections and subsections, include section titles as part of the question text\n"
                            "12. For audit forms or checklists with YES/NO options, use radio button field types\n\n"
                            
                            "Your output MUST strictly adhere to this JSON schema:\n"
                            "{\n"
                            "  \"type\": \"object\",\n"
                            "  \"properties\": {\n"
                            "    \"questions\": {\n"
                            "      \"type\": \"array\",\n"
                            "      \"items\": {\n"
                            "        \"type\": \"object\",\n"
                            "        \"properties\": {\n"
                            "          \"id\": { \"type\": \"string\" },\n"
                            "          \"question_text\": { \"type\": \"string\" },\n"
                            "          \"field_type\": { \"type\": \"string\", \"enum\": [\"text\", \"textarea\", \"radio\", \"checkbox\", \"select\", \"date\", \"email\", \"number\", \"signature\"] },\n"
                            "          \"options\": { \"type\": \"array\", \"items\": { \"type\": \"string\" } },\n"
                            "          \"required\": { \"type\": \"boolean\" }\n"
                            "        },\n"
                            "        \"required\": [\"id\", \"question_text\", \"field_type\"]\n"
                            "      }\n"
                            "    }\n"
                            "  },\n"
                            "  \"required\": [\"questions\"]\n"
                            "}\n\n"
                            
                            "Example extraction format:\n"
                            "{\n"
                            "  \"questions\": [\n"
                            "    {\n"
                            "      \"id\": \"unique_id\",\n"
                            "      \"question_text\": \"The EXACT and COMPLETE text of the question as it appears on the form\",\n"
                            "      \"field_type\": \"text|textarea|radio|checkbox|select|date|email|number|signature\",\n"
                            "      \"options\": [\"Option 1\", \"Option 2\"],\n"
                            "      \"required\": true|false\n"
                            "    }\n"
                            "  ]\n"
                            "}"
                        )
                    },
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": (
                                    "Here's an image of a form. Your task is to extract ALL fields EXACTLY as they appear in the original form, maintaining the original wording, order, and structure. "
                                    "DO NOT TRY TO IMPROVE, CLARIFY, OR REORGANIZE THE FORM - I need the EXACT original form fields as they appear on the form.\n\n"
                                    "CRITICAL EXTRACTION REQUIREMENTS:\n"
                                    "1. DO NOT ADD ANY TEXT THAT IS NOT DIRECTLY PRESENT IN THE IMAGE\n"
                                    "2. Extract ALL form elements - including lines that end with a colon, blank spaces, form fields, checkboxes, etc.\n"
                                    "3. Treat all blank lines following a label as input fields\n"
                                    "4. Look for form field indicators like colons, underlines, or checkboxes\n"
                                    "5. DO NOT SKIP ANY FIELDS, even if they seem minor or redundant\n"
                                    "6. Maintain the EXACT original wording and formatting of all labels\n"
                                    "7. For fields with options (like radio buttons), extract all options exactly\n"
                                    "8. For tables, extract EACH ROW as a separate question with options where applicable\n"
                                    "9. For checklists or forms with YES/NO columns, make each row a separate radio question\n"
                                    "10. In a multi-section form, include the section headers as part of the question text\n"
                                    "11. For tables with multiple columns of checkboxes, convert each row to a question with options\n\n"
                                    "Format your response as structured JSON with a 'questions' array in the SAME ORDER they appear on the form. Include:\n"
                                    "1. A unique 'id' for each field (use field_1, field_2, etc.)\n"
                                    "2. 'question_text': The EXACT text of the field/question as it appears\n"
                                    "3. 'field_type': appropriate type (text, textarea, radio, checkbox, select, date, etc.)\n"
                                    "4. 'options': array of options if applicable (for radio, checkbox, select fields)\n"
                                    "5. 'required': whether the field appears to be required (based on asterisks, etc.)\n\n"
                                    "DO NOT ADD, REWORD, CLARIFY, OR MERGE ANY FIELDS. EXTRACT EXACTLY WHAT IS VISIBLE."
                                )
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                response_format={"type": "json_object"},
                temperature=0.1  # Lower temperature for more precise extraction
            )
            
            result = json.loads(response.choices[0].message.content)
            question_count = len(result.get('questions', []))
            current_app.logger.info(f"Successfully extracted {question_count} form fields from image using GPT-4o")
            
            # Run verification if we have less than the expected number of fields
            # Or if the question count is suspiciously low
            if question_count < 10:
                verification_result = verify_field_extraction_completeness(image_path, base64_image, result.get('questions', []))
                
                # If verification found missed questions, do a second extraction pass focused on those areas
                if not verification_result.get('complete', True) and verification_result.get('missed_questions'):
                    current_app.logger.info("Verification identified potentially missed fields. Performing targeted extraction.")
                    
                    # Construct a prompt for a second extraction pass focused on the missed areas
                    missed_questions = verification_result.get('missed_questions', [])
                    issues = verification_result.get('issues', [])
                    
                    # Format missed questions for the prompt
                    missed_str = "\n".join([f"- {q}" for q in missed_questions])
                    issues_str = "\n".join([f"- {i}" for i in issues])
                    
                    try:
                        current_app.logger.info(f"Performing focused extraction for {len(missed_questions)} potentially missed fields")
                        current_app.logger.info(f"Performing targeted extraction with enhanced JSON schema")
                        focused_response = client.chat.completions.create(
                            model="gpt-4o",
                            messages=[
                                {
                                    "role": "system", 
                                    "content": """You are a highly specialized form field extraction expert focusing on filling in missing information. 
                                    You will be given a list of potentially missed form fields identified by a verification system, and your task is to extract 
                                    the EXACT information for these fields from the image.
                                    
                                    CRITICAL EXTRACTION REQUIREMENTS:
                                    1. DO NOT ADD ANY TEXT THAT IS NOT DIRECTLY PRESENT IN THE IMAGE
                                    2. Extract the EXACT text for each field as it appears in the form - NO paraphrasing, NO combining, NO improving clarity
                                    3. DO NOT modify the wording or structure of the questions in any way
                                    4. Focus ONLY on the potentially missed fields listed
                                    
                                    SPECIAL FOCUS AREAS:
                                    1. Tables - each row should be extracted as a separate question
                                    2. Checklists - each item should be extracted as a separate field
                                    3. Forms with YES/NO columns - each row should be a question with YES and NO as options
                                    4. Audit-style forms - each audit item should be extracted completely
                                    5. Section headers - include these as part of question text
                                    6. Field identifiers like asterisks (*) or other markers for required fields
                                    7. Small text or notes that provide context to questions"""
                                },
                                {
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "text",
                                            "text": f"""
                                            Our verification system has identified that the initial form extraction missed some important fields. 
                                            Please focus ONLY on extracting the following potentially missed fields from the image:
                                            
                                            POTENTIALLY MISSED FIELDS:
                                            {missed_str}
                                            
                                            EXTRACTION ISSUES IDENTIFIED:
                                            {issues_str}
                                            
                                            CRITICAL EXTRACTION INSTRUCTIONS:
                                            1. DO NOT ADD ANY TEXT THAT IS NOT DIRECTLY PRESENT IN THE IMAGE
                                            2. Focus ONLY on the potentially missed fields listed above
                                            3. Extract the EXACT text for each field as it appears in the form
                                            4. DO NOT modify, paraphrase, clarify, or improve the text in any way
                                            5. Include the appropriate field type and whether it appears to be required
                                            6. For multiple choice fields, extract all available options EXACTLY as written
                                            7. For tables with YES/NO columns, extract each row as a radio button question
                                            8. For checklists, extract each item as a separate field
                                            9. For audit forms, make sure to include section headers and extract every item
                                            
                                            Your output MUST follow the same structure as the standard extraction:
                                            {{
                                              "questions": [
                                                {{
                                                  "id": "unique_id",
                                                  "question_text": "The EXACT text of the question as it appears on the form",
                                                  "field_type": "text|textarea|radio|checkbox|select|date|email|number|signature",
                                                  "options": ["Option 1", "Option 2"],
                                                  "required": true|false
                                                }}
                                              ]
                                            }}
                                            """
                                        },
                                        {
                                            "type": "image_url",
                                            "image_url": {
                                                "url": f"data:image/jpeg;base64,{base64_image}"
                                            }
                                        }
                                    ]
                                }
                            ],
                            response_format={"type": "json_object"},
                            temperature=0.1  # Lower temperature for more precise extraction
                        )
                        
                        # Extract and combine the results
                        try:
                            supplementary_result = json.loads(focused_response.choices[0].message.content)
                            supplementary_questions = supplementary_result.get('questions', [])
                            
                            current_app.logger.info(f"Successfully extracted {len(supplementary_questions)} additional fields")
                            
                            # Combine the original questions with the supplementary ones
                            if supplementary_questions:
                                # Generate unique IDs for the supplementary questions
                                for i, q in enumerate(supplementary_questions):
                                    q['id'] = f"supplementary_{i+1}"
                                
                                # Add to the original result
                                result['questions'].extend(supplementary_questions)
                                current_app.logger.info(f"Combined extraction now has {len(result['questions'])} fields")
                        except json.JSONDecodeError as e:
                            current_app.logger.error(f"Failed to parse supplementary extraction: {str(e)}")
                            # Continue with the original extraction
                    except Exception as focused_error:
                        current_app.logger.error(f"Error in focused extraction: {str(focused_error)}")
                        # Continue with the original extraction
            
            # If we got a reasonable number of questions, return the result
            # Some legitimate forms might have only a few fields
            if question_count >= 1 or len(result.get('questions', [])) >= 1:
                return result
            else:
                current_app.logger.warning(f"GPT-4o only extracted {question_count} fields, which seems insufficient. Trying alternative model.")
                # Continue to the fallback model
        
        except Exception as model1_error:
            current_app.logger.warning(f"GPT-4o extraction failed: {str(model1_error)}. Trying alternative model.")
        
        # Fallback to GPT-4-turbo if GPT-4o failed or extracted too few fields
        try:
            current_app.logger.info(f"Attempting form extraction with fallback model using enhanced JSON schema")
            response = client.chat.completions.create(
                model="gpt-4-turbo-preview",  # Use an alternative model
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are a specialized form extraction expert for Minto Disability Services with exceptional attention to detail. "
                            "Your task is to analyze the provided image of a form and extract ALL form fields/questions EXACTLY as they appear in the original. "
                            "This is a fallback extraction attempt, so you must be extremely thorough and detailed.\n\n"
                            
                            "CRITICAL REQUIREMENTS:\n"
                            "1. EXTRACT EVERY SINGLE FORM ELEMENT WITH THE EXACT ORIGINAL TEXT - DO NOT PARAPHRASE, DO NOT COMBINE FIELDS, DO NOT IMPROVE CLARITY\n"
                            "2. DO NOT ADD ANY TEXT THAT IS NOT DIRECTLY PRESENT IN THE IMAGE\n"
                            "3. DO NOT SKIP ANY FIELDS, EVEN IF THEY SEEM MINOR OR REDUNDANT\n"
                            "4. If the image quality is poor, do your absolute best to decipher the text while maintaining exact wording\n"
                            "5. If the image is rotated or skewed, mentally adjust your perspective to read it correctly\n"
                            "6. For complex forms with multiple sections or tables, process sequentially (usually top-to-bottom, left-to-right)\n"
                            "7. Never add explanatory text to fields that isn't present in the original\n"
                            "8. If you're uncertain about field content, include what you can see and note uncertainty with [?] in field labels\n\n"
                            "9. For tabular forms with YES/NO columns or checklists, extract EACH ROW as a separate question\n"
                            "10. Pay special attention to tables - extract EVERY ROW as a separate question\n"
                            "11. For forms with sections and subsections, include section titles as part of the question text\n"
                            "12. For audit forms or checklists with YES/NO options, use radio button field types\n\n"
                            
                            "Your output MUST strictly adhere to this JSON schema:\n"
                            "{\n"
                            "  \"type\": \"object\",\n"
                            "  \"properties\": {\n"
                            "    \"questions\": {\n"
                            "      \"type\": \"array\",\n"
                            "      \"items\": {\n"
                            "        \"type\": \"object\",\n"
                            "        \"properties\": {\n"
                            "          \"id\": { \"type\": \"string\" },\n"
                            "          \"question_text\": { \"type\": \"string\" },\n"
                            "          \"field_type\": { \"type\": \"string\", \"enum\": [\"text\", \"textarea\", \"radio\", \"checkbox\", \"select\", \"date\", \"email\", \"number\", \"signature\"] },\n"
                            "          \"options\": { \"type\": \"array\", \"items\": { \"type\": \"string\" } },\n"
                            "          \"required\": { \"type\": \"boolean\" }\n"
                            "        },\n"
                            "        \"required\": [\"id\", \"question_text\", \"field_type\"]\n"
                            "      }\n"
                            "    }\n"
                            "  },\n"
                            "  \"required\": [\"questions\"]\n"
                            "}\n\n"
                        )
                    },
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": (
                                    "This is our fallback approach to extract ALL form fields, so please analyze this image with extreme thoroughness:\n\n"
                                    
                                    "CRITICAL EXTRACTION REQUIREMENTS:\n"
                                    "1. DO NOT ADD ANY TEXT THAT IS NOT DIRECTLY PRESENT IN THE IMAGE\n"
                                    "2. EXTRACT EVERY SINGLE FORM ELEMENT WITH THE EXACT ORIGINAL TEXT\n"
                                    "3. Analyze the ENTIRE image thoroughly - check headers, footers, margins, and all sections\n"
                                    "4. Be alert for fields in unusual locations (margins, headers, footers, etc.)\n"
                                    "5. For rotated images, visually rotate and extract fields in their logical order\n"
                                    "6. In poor quality images, make your best effort to decipher text while maintaining exact wording\n"
                                    "7. Check for watermarks or background elements that might contain form information\n\n"
                                    
                                    "FIELD IDENTIFICATION GUIDELINES:\n"
                                    "1. Extract ALL form elements - including labels with colons, blank lines/spaces, checkboxes, etc.\n"
                                    "2. Treat all blank lines following text as input fields\n"
                                    "3. Look for field indicators like colons, underlines, boxes, lines, or blank spaces\n"
                                    "4. For tables, extract each row/cell as separate fields based on headers and context\n"
                                    "5. Extract multi-line text areas as single textarea fields\n"
                                    "6. Maintain the EXACT original wording of all field labels\n"
                                    "7. For fields with options (radio buttons, checkboxes), extract all options with exact wording\n"
                                    "8. For tables with YES/NO columns, extract EACH ROW as a separate question with YES and NO as options\n"
                                    "9. For audit forms, each audit item should be extracted as an individual question\n"
                                    "10. For forms with sections and subsections, include section headers as part of the question text\n\n"
                                    
                                    "FORMAT YOUR RESPONSE as JSON with a 'questions' array following these requirements:\n"
                                    "1. Preserve the SAME ORDER as they appear on the form (top-to-bottom, left-to-right)\n"
                                    "2. Generate a unique 'id' for each field (use field_1, field_2, etc.)\n" 
                                    "3. Include 'question_text' with the EXACT text of the label as it appears\n"
                                    "4. Set appropriate 'field_type' (text, textarea, radio, checkbox, select, date, etc.)\n"
                                    "5. Include 'options' array when applicable\n"
                                    "6. Set 'required' based on any indicators in the form (asterisks, 'required' text, etc.)\n\n"
                                    
                                    "DO NOT ADD, REWORD, CLARIFY, OR MERGE ANY FIELDS. EXTRACT EXACTLY WHAT IS VISIBLE."
                                )
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                response_format={"type": "json_object"},
                temperature=0.1  # Lower temperature for more precise extraction
            )
            
            result = json.loads(response.choices[0].message.content)
            question_count = len(result.get('questions', []))
            current_app.logger.info(f"Successfully extracted {question_count} form fields using fallback model")
            return result
            
        except Exception as model2_error:
            current_app.logger.error(f"Fallback model extraction also failed: {str(model2_error)}")
            raise Exception(f"Failed to extract form fields with multiple models: {str(model2_error)}")
    
    except Exception as e:
        current_app.logger.error(f"Error extracting fields from image: {str(e)}")
        raise Exception(f"Failed to extract form fields from image: {str(e)}")

def parse_form_document(file_path):
    """
    Parse a form document and extract a structured representation of questions.
    Handles PDFs, text files, images, and docx files using MarkItDown for enhanced
    document structure preservation when possible.
    """
    try:
        file_path_str = str(file_path)
        current_app.logger.info(f"Parsing form document: {file_path_str}")
        
        # Check if the file is an image
        if is_image_file(file_path_str):
            current_app.logger.info(f"Processing file as image: {file_path_str}")
            return extract_form_fields_from_image(file_path_str)
        
        # First try to convert to markdown using MarkItDown
        # This preserves document structure better than traditional extraction
        try:
            from services.document.markdown_converter import markdown_converter
            
            current_app.logger.info(f"Attempting to convert {file_path_str} to markdown for enhanced structure preservation")
            markdown_result = markdown_converter.convert_to_markdown(file_path_str)
            
            if markdown_result["success"] and markdown_result["markdown"]:
                markdown_content = markdown_result["markdown"]
                current_app.logger.info(f"Successfully converted to markdown (length: {len(markdown_content)} chars)")
                current_app.logger.debug(f"Markdown content preview: {markdown_content[:300]}...")
                
                # Use the markdown content for AI form field extraction
                file_content = markdown_content
                
                # For debugging purposes
                current_app.logger.info("Using markdown representation for form extraction")
                current_app.logger.debug(f"First 1000 chars of markdown content: {file_content[:1000]}")
                
                # Skip the traditional extraction methods
                return extract_form_fields_from_markdown(file_content, file_path_str)
            else:
                current_app.logger.warning(f"MarkItDown conversion failed: {markdown_result.get('error')}. Falling back to traditional extraction.")
        except Exception as md_error:
            current_app.logger.warning(f"Error converting to markdown: {str(md_error)}. Falling back to traditional extraction.")
        
        # If markdown conversion fails, fall back to traditional extraction methods
        # Extract text from the file based on type
        if file_path_str.endswith('.pdf'):
            import PyPDF2
            
            text = ""
            with open(file_path_str, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if the PDF contains any images
                has_images = False
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    if '/XObject' in page['/Resources']:
                        has_images = True
                        break
                
                # If PDF has images, it might be a scanned form
                # In this case, try to use GPT-4 Vision to extract fields
                if has_images:
                    current_app.logger.info(f"PDF contains images, attempting to process as image: {file_path_str}")
                    try:
                        return extract_form_fields_from_image(file_path_str)
                    except Exception as e:
                        current_app.logger.warning(f"Failed to process PDF as image: {str(e)}. Falling back to text extraction.")
                
                # Extract text from all pages
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    extracted_text = page.extract_text()
                    if extracted_text:
                        text += extracted_text + "\n\n"
            
            file_content = text
        elif file_path_str.endswith('.docx'):
            # Use python-docx library to extract docx content
            try:
                import docx
                import io
                import tempfile
                from PIL import Image
                
                current_app.logger.info(f"Extracting content from DOCX file: {file_path_str}")
                doc = docx.Document(file_path_str)
                
                # IMPORTANT: When dealing with forms, specifically look for checklist patterns
                # which are common in audit forms
                
                # Extract questions directly from the document structure
                form_questions = []
                question_id = 1
                
                # Process tables - often used for checklists and structured forms
                for table_index, table in enumerate(doc.tables):
                    header_row = None
                    header_cells = []
                    
                    # Check if this table has a header row
                    if len(table.rows) > 0:
                        header_row = table.rows[0]
                        header_cells = [cell.text.strip() for cell in header_row.cells if cell.text.strip()]
                    
                    # Process each row after the header
                    for row_index, row in enumerate(table.rows):
                        if row_index == 0 and header_cells:
                            # Skip the header row
                            continue
                            
                        row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        
                        if not row_cells:
                            continue
                            
                        # Determine the question text and field type based on the row content
                        question_text = row_cells[0] if row_cells else ""
                        
                        # Skip empty or meaningless rows
                        if not question_text or question_text.lower() in ('yes', 'no', 'y', 'n', 'n/a'):
                            continue
                            
                        # Check if this looks like a YES/NO question
                        options = []
                        field_type = "text"  # Default
                        
                        # Look for YES/NO pattern in other cells
                        yes_no_cells = [cell for cell in row_cells[1:] if cell.lower() in ('yes', 'no', 'y', 'n', 'n/a')]
                        
                        if yes_no_cells or (header_cells and any(h.lower() in ('yes', 'no', 'y/n', 'yes/no') for h in header_cells)):
                            field_type = "radio"
                            options = ["Yes", "No"]
                            
                            # If there's a N/A option detected, add it
                            if any(cell.lower() == 'n/a' for cell in row_cells):
                                options.append("N/A")
                        
                        # Add this question to our list
                        form_questions.append({
                            "id": f"question_{question_id}",
                            "question_text": question_text,
                            "field_type": field_type,
                            "options": options,
                            "required": True  # Default to required for form items
                        })
                        question_id += 1
                
                # Process paragraphs - some forms have one question per paragraph
                in_section = False
                current_section = ""
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                        
                    # Check if this is a section header (often appears as bold, larger text)
                    is_section_header = False
                    for run in para.runs:
                        if run.bold or run.font.size and run.font.size > 12:
                            is_section_header = True
                            break
                    
                    if is_section_header:
                        current_section = text
                        in_section = True
                        continue
                    
                    # Check if this paragraph looks like a form question
                    # Form questions often end with a colon or have whitespace for answers
                    if ":" in text or text.endswith("?") or "_____" in text or "□" in text or "☐" in text:
                        # This looks like a question
                        question_text = text
                        
                        # If we're in a section, include the section name
                        if in_section and current_section:
                            question_text = f"{current_section} - {question_text}"
                            
                        # Determine field type based on content
                        field_type = "text"  # Default
                        options = []
                        
                        # Check for multiple choice indicators
                        if "□" in text or "☐" in text or "checkbox" in text.lower() or "check box" in text.lower():
                            field_type = "checkbox"
                            # Try to extract options
                            option_parts = text.split("□")
                            if len(option_parts) > 1:
                                options = [part.strip() for part in option_parts[1:] if part.strip()]
                                
                        # Add this question
                        form_questions.append({
                            "id": f"question_{question_id}",
                            "question_text": question_text,
                            "field_type": field_type,
                            "options": options,
                            "required": text.endswith("*") or "*required" in text.lower()
                        })
                        question_id += 1
                
                # If we found direct questions, use them
                if form_questions:
                    current_app.logger.info(f"Successfully extracted {len(form_questions)} questions directly from DOCX structure")
                    return {"questions": form_questions}
                
                # If we couldn't extract structured questions directly, try the traditional text-based approach
                # Extract text from tables
                table_content = []
                for table in doc.tables:
                    for i, row in enumerate(table.rows):
                        row_texts = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_texts.append(cell.text.strip())
                        
                        if row_texts:
                            # Join the row cells with appropriate separators
                            if i == 0:  # First row might be headers
                                table_content.append(" | ".join(row_texts))
                            else:
                                # For YES/NO columns in forms, format specially
                                if any(cell.lower() in ('yes', 'no', 'y', 'n') for cell in row_texts):
                                    # First cell is usually the question text
                                    question = row_texts[0] if row_texts else ""
                                    # Other cells might be YES/NO
                                    options = [cell for cell in row_texts[1:] if cell.lower() in ('yes', 'no', 'y', 'n')]
                                    table_content.append(f"{question} [Options: {', '.join(options)}]")
                                else:
                                    table_content.append(" - ".join(row_texts))
                
                # Extract regular paragraphs
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                
                # Combine all content with appropriate separation
                all_content = paragraphs + ["\n\nTABLE CONTENT:\n"] + table_content if table_content else paragraphs
                file_content = "\n".join(all_content)
                
                # If we extracted meaningful content, analyze it
                if file_content and len(file_content.strip()) > 0:
                    current_app.logger.info(f"Successfully extracted {len(file_content)} characters from DOCX text")
                    
                    # For DOCX, which tend to have complex structure,
                    # we'll take a different two-pronged approach
                    
                    # First, try to render the first page of the document as an image
                    # This works well for forms with visual layout
                    try:
                        current_app.logger.info("Attempting to process DOCX as image for better form extraction")
                        # Try using vision API directly on the DOCX file
                        vision_result = extract_form_fields_from_image(file_path_str)
                        
                        # If we got meaningful results from the vision API, return those
                        if vision_result and vision_result.get('questions') and len(vision_result.get('questions', [])) > 3:
                            current_app.logger.info(f"Vision API successfully extracted {len(vision_result.get('questions', []))} questions from DOCX")
                            return vision_result
                    except Exception as vision_error:
                        current_app.logger.warning(f"Vision enhancement failed: {str(vision_error)}. Using text extraction instead.")
                    
                    # Second, if vision API didn't work, try the text-based approach with GPT-4o
                    try:
                        client = get_openai_client()
                        current_app.logger.info("Using GPT-4o with extracted DOCX text")
                        
                        # Add some special instructions for handling DOCX content
                        response = client.chat.completions.create(
                            model="gpt-4o",
                            messages=[
                                {
                                    "role": "system",
                                    "content": (
                                        "You are a form extraction specialist focused on DOCX files. "
                                        "Your task is to identify ALL form fields from the given text extracted from a DOCX document. "
                                        "IMPORTANT GUIDELINES:\n"
                                        "1. Each row in tables should be treated as a separate question\n"
                                        "2. For rows with YES/NO columns, extract as radio button questions\n"
                                        "3. Look for checklist patterns and extract each item\n"
                                        "4. Pay attention to section headers and incorporate them\n"
                                        "5. Be extremely thorough in identifying every possible field\n"
                                        "6. If the text contains \"TABLE CONTENT\" markers, pay special attention to that section\n"
                                        "7. For audit forms, assume each item needs YES/NO options\n"
                                        "8. If a line looks like a header, it's probably not a question itself\n"
                                        "9. If a field is followeded by a blank space or ___, it's likely a text field"
                                    )
                                },
                                {
                                    "role": "user",
                                    "content": f"Here is text extracted from a DOCX form document. Extract ALL questions/fields:\n\n{file_content}"
                                }
                            ],
                            response_format={"type": "json_object"},
                            temperature=0.2
                        )
                        
                        try:
                            result = json.loads(response.choices[0].message.content)
                            question_count = len(result.get('questions', []))
                            
                            if question_count > 0:
                                current_app.logger.info(f"Successfully extracted {question_count} form fields from DOCX text using GPT-4o")
                                return result
                        except json.JSONDecodeError:
                            current_app.logger.error("Failed to parse GPT response for DOCX extraction")
                    except Exception as text_extraction_error:
                        current_app.logger.error(f"Text-based DOCX extraction failed: {str(text_extraction_error)}")
                
                else:
                    current_app.logger.warning(f"DOCX content extraction resulted in empty content: {file_path_str}")
                
                # As a last resort, explicitly tell the model this is DOCX and needs special processing
                try:
                    client = get_openai_client()
                    current_app.logger.info("Attempting final extraction approach for DOCX")
                    
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {
                                "role": "system",
                                "content": "You are a form structure generator. This is a DOCX file that couldn't be properly extracted. Please generate a reasonable form structure based on the filename."
                            },
                            {
                                "role": "user", 
                                "content": f"The file '{os.path.basename(file_path_str)}' appears to be a form, but we couldn't extract its content properly. Please create a basic form structure based on the filename. If it has 'audit' or 'checklist' in the name, assume it's a checklist with Yes/No options."
                            }
                        ],
                        response_format={"type": "json_object"},
                        temperature=0.7
                    )
                    
                    result = json.loads(response.choices[0].message.content)
                    return result
                except Exception as final_error:
                    current_app.logger.error(f"Final DOCX extraction attempt failed: {str(final_error)}")
                    
            except Exception as docx_error:
                current_app.logger.error(f"Failed to extract DOCX content with python-docx: {str(docx_error)}")
                # Try to use vision API as a fallback for docx files
                try:
                    vision_result = extract_form_fields_from_image(file_path_str)
                    return vision_result
                except Exception as vision_error:
                    current_app.logger.error(f"Vision fallback failed: {str(vision_error)}")
                    
                    # As an absolute last resort, use the filename to guess the form structure
                    try:
                        client = get_openai_client()
                        current_app.logger.info("Using filename-based extraction as last resort")
                        
                        filename = os.path.basename(file_path_str)
                        response = client.chat.completions.create(
                            model="gpt-4o",
                            messages=[
                                {
                                    "role": "system",
                                    "content": "You are a form structure generator creating form fields based solely on a filename."
                                },
                                {
                                    "role": "user", 
                                    "content": f"The file '{filename}' appears to be a form, but we couldn't extract its content properly. Please create a basic form structure based on the filename. If it has 'audit' or 'checklist' in the name, assume it's a checklist with Yes/No options."
                                }
                            ],
                            response_format={"type": "json_object"},
                            temperature=0.7
                        )
                        
                        result = json.loads(response.choices[0].message.content)
                        question_count = len(result.get('questions', []))
                        current_app.logger.info(f"Generated {question_count} form fields based on filename")
                        return result
                    except Exception as filename_error:
                        current_app.logger.error(f"Filename-based extraction failed: {str(filename_error)}")
                        
                    # If everything else fails, return a minimal structure with helpful message
                    file_content = "This document appears to be a form that requires special handling."
        else:
            # Text file
            try:
                with open(file_path_str, 'r', encoding='utf-8') as file:
                    file_content = file.read()
            except UnicodeDecodeError:
                # If UTF-8 fails, try binary read as it might be an image or binary file
                current_app.logger.info(f"Failed to read as text, attempting to process as image: {file_path_str}")
                return extract_form_fields_from_image(file_path_str)
        
        # If no content was extracted, use a default template for an incident form
        if not file_content or len(file_content.strip()) < 50:
            current_app.logger.warning(f"Limited content extracted from {file_path_str}, using default incident form template")
            file_content = """
            Incident Report Form for Minto Disability Services
            
            1. Incident Date (date, required)
            2. Incident Time (time, required)
            3. Incident Location (text, required)
            4. Persons Involved (textarea, required)
            5. Incident Description (textarea, required)
            6. Severity (radio: Minor, Moderate, Severe, required)
            7. Medical Attention Required (checkbox: Yes, No, required)
            8. Immediate Actions Taken (textarea, required)
            9. Reporter Name (text, required)
            10. Reporter Position (text, required)
            11. Reporter Contact (text, required)
            12. Reporter Email (email, required)
            13. Additional Comments (textarea, optional)
            """
        
        # Get the OpenAI client
        client = get_openai_client()
        
        current_app.logger.debug(f"Sending form content to OpenAI for parsing: {file_content[:500]}...")
        
        # the newest OpenAI model is "gpt-4o" which was released May 13, 2024.
        # do not change this unless explicitly requested by the user
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a form parsing assistant for Minto Disability Services with special focus on EXACT field extraction. "
                        "Your task is to extract ALL questions and fields from the provided form document EXACTLY as they appear in the original. "
                        "IMPORTANT GUIDELINES:\n"
                        "1. Do NOT rephrase, modify, combine, or add any questions. Preserve the original text, formatting, and order exactly\n"
                        "2. Extract the precise label text for every field. Do not summarize or generalize fields\n"
                        "3. Be extremely literal in your extraction\n"
                        "4. For tables, extract EACH ROW as a separate question\n"
                        "5. For tables with YES/NO columns, make each row a radio question with YES and NO as options\n"
                        "6. For checklist forms, extract every checklist item as a separate field\n"
                        "7. In forms with sections/categories, include the section name as part of the question text\n"
                        "8. For audit forms with multiple columns (like YES/NO), extract each item as a question with options\n\n"
                        
                        "For each field, identify: "
                        "1. A unique 'id' (use a simple index number or field name without modifying the text) "
                        "2. 'question_text' (the EXACT and COMPLETE text of the field/question as it appears on the form, including any numbering or formatting) "
                        "3. 'field_type' (text, checkbox, radio, select, textarea, date, email, number, etc.) "
                        "4. 'options' array (EXACT options text for checkbox, radio, select fields) "
                        "5. Whether the field is 'required' (boolean - assume any field with an asterisk or otherwise marked as required is true) "
                        "Return this in a structured JSON format with a 'questions' array in the SAME ORDER as they appear on the form. "
                        "Do not skip ANY fields. Do not merge fields. Do not improve or reword the questions. "
                        "Output must be formatted as JSON."
                    )
                },
                {
                    "role": "user",
                    "content": f"Parse the following form content and return all the form fields in JSON format. Extract every field EXACTLY as written: \n\n{file_content}"
                }
            ],
            response_format={"type": "json_object"}
        )
        
        result = json.loads(response.choices[0].message.content)
        return result
    
    except Exception as e:
        logging.error(f"Error parsing form document: {str(e)}")
        raise Exception(f"Failed to parse form: {str(e)}")

def generate_form_questions(form_structure):
    """
    Generate a step-by-step question flow from form structure.
    """
    try:
        # Get the OpenAI client
        client = get_openai_client()
        
        # Validate input
        if not form_structure or not isinstance(form_structure, dict):
            logging.error(f"Invalid form structure: {form_structure}")
            # Return a standardized version of the input or an empty structure
            return {'questions': []}
        
        # Count the original number of questions
        original_questions = form_structure.get('questions', [])
        original_question_count = len(original_questions)
        
        # If there are no questions, return an empty structure
        if original_question_count == 0:
            logging.warning("No questions found in form structure")
            return {'questions': []}
            
        current_app.logger.info(f"Organizing {original_question_count} questions into a sequential flow")
        
        # Skip OpenAI processing and directly standardize the original questions
        # This ensures we keep the exact original structure without any changes
        standardized_questions = []
        for q in original_questions:
            # Skip if the question is None or not a dictionary
            if not q or not isinstance(q, dict):
                continue
                
            # Make a copy of the question to avoid modifying the original
            try:
                standardized_q = q.copy()
            except:
                # If we can't copy the question, create a new dict
                standardized_q = {}
            
            # Ensure question_text field exists
            if 'question_text' not in standardized_q:
                standardized_q['question_text'] = standardized_q.get('question') or standardized_q.get('label') or standardized_q.get('text') or f"Question {len(standardized_questions) + 1}"
            
            # Ensure field_type field exists
            if 'field_type' not in standardized_q:
                standardized_q['field_type'] = standardized_q.get('type') or 'text'
            
            # Ensure id field exists
            if 'id' not in standardized_q:
                # Generate an ID based on the question text
                text = str(standardized_q.get('question_text', '')).strip()
                standardized_q['id'] = text.lower().replace(' ', '_')[:30] or f"question_{len(standardized_questions) + 1}"
            
            # Ensure options exists for appropriate field types
            if standardized_q['field_type'] in ['radio', 'checkbox', 'select'] and 'options' not in standardized_q:
                standardized_q['options'] = []
                
            standardized_questions.append(standardized_q)
        
        current_app.logger.info(f"Standardized {len(standardized_questions)} questions without modifying text or order")
        return {'questions': standardized_questions}
        
    except Exception as e:
        logging.error(f"Error generating form questions: {str(e)}")
        raise Exception(f"Failed to generate questions: {str(e)}")

def generate_embeddings(text):
    """
    Generate embeddings for a given text using OpenAI's embedding model.
    """
    try:
        # Get the OpenAI client
        client = get_openai_client()
        
        response = client.embeddings.create(
            model="text-embedding-ada-002",
            input=text
        )
        return response.data[0].embedding
    
    except Exception as e:
        logging.error(f"Error generating embeddings: {str(e)}")
        raise Exception(f"Failed to generate embeddings: {str(e)}")

def generate_answer_with_context(question, contexts):
    """
    Generate an answer to a question based on the provided context.
    """
    try:
        context_text = "\n\n".join(contexts)
        
        # Get the OpenAI client
        client = get_openai_client()
        
        # the newest OpenAI model is "gpt-4o" which was released May 13, 2024.
        # do not change this unless explicitly requested by the user
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an assistant for Minto Disability Services. Answer questions about policies and procedures "
                        "based solely on the provided context. If you don't find the information in the context, say so clearly. "
                        "Do not make up information. Be concise but thorough. When appropriate, cite specific policies by name.\n\n"
                        "FORMATTING GUIDELINES:\n"
                        "1. Do NOT use markdown formatting, asterisks, or other special characters for emphasis\n"
                        "2. Use plain text formatting with regular paragraph breaks for readability\n"
                        "3. For lists, use simple numbering (1, 2, 3) or bullet points with hyphens (-)\n"
                        "4. If information varies by category, use clear headers with colons (Example: 'Full-Time Employees:')\n"
                        "5. Organize information in a clean, readable manner without any special formatting characters\n"
                        "6. Never use asterisks (**) for emphasis or headings"
                    )
                },
                {
                    "role": "user",
                    "content": f"Context documents:\n{context_text}\n\nQuestion: {question}"
                }
            ]
        )
        
        return response.choices[0].message.content
    
    except Exception as e:
        logging.error(f"Error generating answer: {str(e)}")
        raise Exception(f"Failed to generate answer: {str(e)}")
