"""
Template-based fallback solution for the Home Safety Checklist
This module provides a predefined field structure for the specific form
to ensure consistent extraction even when AI-based extraction fails.
"""

from typing import List, Dict, Any


def get_home_safety_checklist_template() -> List[Dict[str, Any]]:
    """
    Returns a predefined template for the Home Safety Checklist
    with the exact structure matching the document.
    """
    return [
        {
            "id": "participant_name",
            "question_text": "Participant Name:",
            "field_type": "text",
            "required": True
        },
        {
            "id": "address",
            "question_text": "Address:",
            "field_type": "textarea",
            "required": True
        },
        {
            "id": "date_of_assessment",
            "question_text": "Date of assessment:",
            "field_type": "date",
            "required": True
        },
        # ENTRANCE TO HOME section
        {
            "id": "entrance_outside_lights",
            "question_text": "Are there outside lights covering the sidewalks and/or other entrance ways?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "entrance_steps_sidewalks",
            "question_text": "Are the steps & sidewalks in good repair and free from debris/material?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "entrance_ramp_needed",
            "question_text": "Is a ramp needed?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "entrance_railings_secured",
            "question_text": "Are the railings on the steps secured?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "entrance_peephole",
            "question_text": "Is there a functional peephole in the front door?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "entrance_deadbolt",
            "question_text": "Does the door have a deadbolt lock that does not require a key to open it from the inside (unless client tends to wander)?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        # GENERAL section
        {
            "id": "general_emergency_plan",
            "question_text": "Is there an emergency plan in place?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_smoke_detectors",
            "question_text": "Are working smoke detectors installed?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_fire_extinguisher",
            "question_text": "Is there a \"ready-to-use\" fire extinguisher(s) on the premises?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_clutter_free",
            "question_text": "Are inside halls and stairways free of clutter/debris?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_throw_rugs",
            "question_text": "Are throw rugs removed?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_handrails",
            "question_text": "Are there sturdy handrails or banisters by all steps and stairs?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_electrical_cords",
            "question_text": "Are electrical cords non-frayed and placed in a manner to avoid tripping?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_outlets_overloaded",
            "question_text": "Are electric outlets/switches overloaded (e.g., warm to the touch)?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_rugs_secured",
            "question_text": "Are rugs secured around the edges?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_hazardous_products",
            "question_text": "Are hazardous products labelled and kept in a secure place?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_stool_needed",
            "question_text": "Is there a need for a stool to reach high shelves/cupboards?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_smoking_safety",
            "question_text": "Is smoking paraphernalia handled safely (e.g., cigarettes put out)?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_smoking_oxygen",
            "question_text": "Does anybody smoke in homes where oxygen is in use?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_pets_controlled",
            "question_text": "Are all animals/pets, on site, controlled?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_free_from_pests",
            "question_text": "Is the home free from bugs, mice and/or animal waste?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_materials_stored",
            "question_text": "Are materials stored safely and at a proper height?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_emergency_response",
            "question_text": "Does the client wear an emergency response necklace/bracelet?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_polished_floors",
            "question_text": "Are polished floors no waxed or waxed-free?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "general_weapons",
            "question_text": "Are there any weapon on the premises?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        # MEDICATIONS section
        {
            "id": "medications_clearly_marked",
            "question_text": "Are all medications marked clearly?",
            "field_type": "radio",
            "options": ["Yes", "No", "N/A"],
            "required": True
        },
        {
            "id": "medications_named",
            "question_text": "Are medications named?",
            "field_type": "radio",
            "options": ["Yes", "No", "N/A"],
            "required": True
        },
        {
            "id": "medications_dated",
            "question_text": "Are medications dated?",
            "field_type": "radio",
            "options": ["Yes", "No", "N/A"],
            "required": True
        },
        {
            "id": "medications_instructions",
            "question_text": "Are instructions given as to how medications are to be taken?",
            "field_type": "radio",
            "options": ["Yes", "No", "N/A"],
            "required": True
        },
        {
            "id": "medications_disposed",
            "question_text": "Are out-of-date medications disposed of safely?",
            "field_type": "radio",
            "options": ["Yes", "No", "N/A"],
            "required": True
        },
        {
            "id": "medications_locked",
            "question_text": "Are medications locked away (especially where children are present or the client is forgetful)?",
            "field_type": "radio",
            "options": ["Yes", "No", "N/A"],
            "required": True
        },
        {
            "id": "medications_blister_packs",
            "question_text": "Are blister packs available?",
            "field_type": "radio",
            "options": ["Yes", "No", "N/A"],
            "required": True
        },
        # BATHROOM section
        {
            "id": "bathroom_grab_bars",
            "question_text": "Are there grab bars installed beside the toilet and in the bath or shower?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "bathroom_bath_seat",
            "question_text": "Is a bath seat or bath board required?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "bathroom_commode",
            "question_text": "Is a raised toilet seat and/or commode required?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "bathroom_nonslip_surfaces",
            "question_text": "Are nonslip surfaces or adhesive strips in the bath or shower?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "bathroom_hand_held_shower",
            "question_text": "Is a hand-held shower head required?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "bathroom_hot_water",
            "question_text": "Is hot water set at safe temperature to prevent scalding?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "bathroom_ventilation_fan",
            "question_text": "Is ventilation fan used or is window open during showers?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        # FURNITURE section
        {
            "id": "furniture_castors",
            "question_text": "Does furniture have castors or roll?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "furniture_chairs_sit_stand",
            "question_text": "Are chairs appropriate for sit to stand?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "furniture_bed_height",
            "question_text": "Is bed height appropriate?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "furniture_safety_devices",
            "question_text": "Are safety devices recommended for the bed (bed levers, bed stick, leg raisers)?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        # KITCHEN section
        {
            "id": "kitchen_frequently_used_accessible",
            "question_text": "Are frequently used items accessible without climbing?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "kitchen_stove_knobs",
            "question_text": "Are stove knobs clearly marked?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "kitchen_handles_pots",
            "question_text": "Are handles on pots turned towards the centre of the stove?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "kitchen_eating_location",
            "question_text": "Is there an appropriate place to eat?",
            "field_type": "radio",
            "options": ["Yes", "No"],
            "required": True
        },
        {
            "id": "safety_concerns",
            "question_text": "Safety Concerns and Comments:",
            "field_type": "textarea",
            "required": False
        },
        {
            "id": "priority_concerns",
            "question_text": "List priority safety concerns to be addressed:",
            "field_type": "textarea",
            "required": False
        },
        {
            "id": "assessor_name",
            "question_text": "Name of Assessor:",
            "field_type": "text",
            "required": True
        },
        {
            "id": "assessor_signature",
            "question_text": "Signature of Assessor:",
            "field_type": "signature",
            "required": True
        },
        {
            "id": "signature_date",
            "question_text": "Date:",
            "field_type": "date",
            "required": True
        }
    ]


def is_home_safety_checklist(file_path: str) -> bool:
    """
    Determine if the file is likely a Home Safety Checklist.
    
    Args:
        file_path: The path to the file
        
    Returns:
        bool: True if the file appears to be a home safety checklist
    """
    # This should be updated with more reliable methods of identification
    try:
        # If it's a file path, read content
        if isinstance(file_path, str) and os.path.exists(file_path):
            try:
                import docx
                doc = docx.Document(file_path)
                content = ""
                for para in doc.paragraphs:
                    content += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                content += paragraph.text + "\n"
                file_path = content
            except:
                # If we can't read as a docx, try to read as text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    file_path = f.read()
        
        # Check for key phrases that indicate it's a home safety checklist
        lower_content = file_path.lower()
        
        # Define key markers that strongly indicate this form
        key_markers = [
            "home safety checklist",
            "safety criteria",
            "entrance to home",
            "are there outside lights",
            "are the steps & sidewalks"
        ]
        
        # Count how many key markers are found
        marker_count = sum(1 for marker in key_markers if marker in lower_content)
        
        # If at least 3 key markers are found, it's likely this form
        return marker_count >= 3
        
    except Exception as e:
        print(f"Error checking if document is a home safety checklist: {str(e)}")
        return False


# Add the os import at the top
import os