"""
Template-based fallback solution for the Root Cause Analysis Form
This module provides a predefined field structure for the specific form
to ensure consistent extraction even when AI-based extraction fails.
"""

from typing import List, Dict, Any
import re
import os
import docx

def get_root_cause_analysis_template() -> List[Dict[str, Any]]:
    """
    Returns a predefined template for the Root Cause Analysis Form
    with the exact structure matching the document.
    """
    template = []
    
    # Participant and Event Information section
    template.append({
        "id": "q_participant_file_id",
        "question_text": "Participant File ID:",
        "field_type": "text",
        "required": True
    })
    
    template.append({
        "id": "q_participant_dob",
        "question_text": "Participant DOB:",
        "field_type": "date",
        "required": True
    })
    
    template.append({
        "id": "q_program_facility",
        "question_text": "Program/Facility:",
        "field_type": "text",
        "required": True
    })
    
    template.append({
        "id": "q_gender",
        "question_text": "Gender:",
        "field_type": "text",
        "required": True
    })
    
    template.append({
        "id": "q_location_of_event",
        "question_text": "Location of event:",
        "field_type": "text",
        "required": True
    })
    
    template.append({
        "id": "q_date_of_event",
        "question_text": "Date of Event:",
        "field_type": "date",
        "required": True
    })
    
    template.append({
        "id": "q_date_rca_completed",
        "question_text": "Date RCA Completed:",
        "field_type": "date",
        "required": True
    })
    
    template.append({
        "id": "q_staff_participated",
        "question_text": "Name of staff participated in RCA:",
        "field_type": "text",
        "required": True
    })
    
    template.append({
        "id": "q_rca_team_leader",
        "question_text": "RCA team leader name:",
        "field_type": "text",
        "required": True
    })
    
    template.append({
        "id": "q_rca_team_members",
        "question_text": "RCA Team Members:",
        "field_type": "textarea",
        "required": True
    })
    
    # Event Description section
    template.append({
        "id": "q_event_description",
        "question_text": "1. THE EVENT – Describe what happened and any harm that resulted. Identify the proximate cause, if known.",
        "field_type": "textarea",
        "required": True
    })
    
    # Background & Factors section
    template.append({
        "id": "q_expected_sequence",
        "question_text": "2.1 What was the sequence of events that was expected to take place?",
        "field_type": "textarea",
        "required": True
    })
    
    template.append({
        "id": "q_deviation_from_expected",
        "question_text": "2.2 Was there a deviation from the expected sequence?",
        "field_type": "radio",
        "options": ["Yes", "No"],
        "required": True
    })
    
    template.append({
        "id": "q_deviation_description",
        "question_text": "If YES, describe the deviation.",
        "field_type": "textarea",
        "required": False
    })
    
    template.append({
        "id": "q_deviation_contributed",
        "question_text": "2.3 Was any deviation from the expected sequence likely to have led to or contributed to the adverse event?",
        "field_type": "radio",
        "options": ["Yes", "No", "NA"],
        "required": True
    })
    
    template.append({
        "id": "q_causal_statement",
        "question_text": "If YES, describe with causal statement.",
        "field_type": "textarea",
        "required": False
    })
    
    template.append({
        "id": "q_sequence_in_policy",
        "question_text": "2.4 Was the expected sequence described in policy, procedure, written guidelines, or included in staff training?",
        "field_type": "radio",
        "options": ["Yes", "No", "NA"],
        "required": True
    })
    
    template.append({
        "id": "q_policy_source",
        "question_text": "If YES, cite source.",
        "field_type": "text",
        "required": False
    })
    
    template.append({
        "id": "q_meets_requirements",
        "question_text": "2.5 Does the expected sequence or process meet regulatory requirements and/or practice standards?",
        "field_type": "radio",
        "options": ["Yes", "No", "NA"],
        "required": True
    })
    
    template.append({
        "id": "q_requirements_deviation",
        "question_text": "If NO, describe deviation from requirements/standards.",
        "field_type": "textarea",
        "required": False
    })
    
    template.append({
        "id": "q_staff_credentialed",
        "question_text": "2.10 Were involved staff credentialed/skilled to carry out the tasks expected of them?",
        "field_type": "radio",
        "options": ["Yes", "No", "NA"],
        "required": True
    })
    
    template.append({
        "id": "q_inadequacy_description",
        "question_text": "If NO, describe the perceived inadequacy.",
        "field_type": "textarea",
        "required": False
    })
    
    template.append({
        "id": "q_communication_issues",
        "question_text": "2.15 Did a lack of communication or incomplete communication contribute to or cause the adverse event?",
        "field_type": "radio",
        "options": ["Yes", "No", "NA"],
        "required": True
    })
    
    template.append({
        "id": "q_communication_contribution",
        "question_text": "If YES, describe who and what and how it contributed.",
        "field_type": "textarea",
        "required": False
    })
    
    template.append({
        "id": "q_rank_factors",
        "question_text": "2.20 Rank order the factors considered responsible for the adverse event, beginning with the proximate cause, followed by the most important to less important contributory factors. Attach Contributory Factors Diagram, if available.",
        "field_type": "textarea",
        "required": True
    })
    
    # Prevention Strategies section
    template.append({
        "id": "q_prevention_strategies",
        "question_text": "4. PREVENTION STRATEGIES – List from highest priority to lowest priority the recommended actions designed to prevent a future occurrence of the adverse event. Begin with a rank of 1 (highest). For each strategy or action provide an estimated cost, if known, and any additional considerations or recommendations for implementing the strategy (e.g., phase-in, immediate need, triage by risk).",
        "field_type": "textarea",
        "required": True
    })
    
    # Recipients table
    template.append({
        "id": "q_recipients_table",
        "question_text": "Forward this report to all RCA team members and to the following individuals:",
        "field_type": "table",
        "columns": ["Name", "Title", "Organization", "Address", "Email"],
        "rows": 3,
        "required": False
    })
    
    return template


def is_root_cause_analysis(file_path_or_content: str) -> bool:
    """
    Determine if the file is likely a Root Cause Analysis Form.
    
    Args:
        file_path_or_content: The path to the file or the content as a string
        
    Returns:
        bool: True if the file appears to be a root cause analysis form
    """
    try:
        content = ""
        
        # Check if the input is a file path
        if os.path.exists(file_path_or_content):
            # Extract text based on file type
            if file_path_or_content.lower().endswith('.docx'):
                doc = docx.Document(file_path_or_content)
                content = '\n'.join([para.text for para in doc.paragraphs])
                
                # Add table text
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            content += f"\n{cell.text}"
            else:
                # Assume it's a text file or already content
                with open(file_path_or_content, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
        else:
            # Assume it's already content
            content = file_path_or_content
            
        # Define keywords and patterns that indicate this is a root cause analysis form
        keywords = [
            r"root\s+cause\s+analysis",
            r"rca\s+team\s+leader",
            r"proximate\s+cause",
            r"contributory\s+factors",
            r"prevention\s+strategies",
            r"sequence\s+of\s+events",
            r"adverse\s+event"
        ]
        
        # Create a scoring system - if enough keywords are found, consider it a match
        score = 0
        for keyword in keywords:
            if re.search(keyword, content.lower()):
                score += 1
                
        # If 3 or more keywords match, consider it a root cause analysis form
        return score >= 3
        
    except Exception:
        # If there's any error in processing, default to False
        return False