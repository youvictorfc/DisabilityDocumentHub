"""
Template-based fallback solution for the Waste Risk Assessment Checklist
This module provides a predefined field structure for the specific form
to ensure consistent extraction even when AI-based extraction fails.
"""

from typing import List, Dict, Any
import os
import re


def get_waste_risk_assessment_template() -> List[Dict[str, Any]]:
    """
    Returns a predefined template for the Waste Risk Assessment Checklist
    with the exact structure matching the document.
    """
    return [
        {
            "id": "assessor_name",
            "question_text": "Name of assessor:",
            "field_type": "text",
            "required": True,
            "options": []
        },
        {
            "id": "assessment_date",
            "question_text": "Date of assessment:",
            "field_type": "date",
            "required": True,
            "options": []
        },
        {
            "id": "location",
            "question_text": "Location:",
            "field_type": "text",
            "required": True,
            "options": []
        },
        {
            "id": "remarks_actions",
            "question_text": "Remarks/Actions:",
            "field_type": "textarea",
            "required": False,
            "options": []
        },
        # General Waste Handling Requirements
        {
            "id": "waste_segregation_posters",
            "question_text": "1. Are waste segregation posters available and displayed in all relevant areas e.g., dirty utility room/waste holding room, clean utility room?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "bins_labeled",
            "question_text": "2. Are all bins labelled as appropriate e.g., healthcare risk waste or clinical waste; healthcare non risk waste or general waste; paper waste; glass waste etc.?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "correct_color_liners",
            "question_text": "3. Have all bins the correct colour coded liner / bag e.g., yellow bag for clinical waste and black or clear bag for non-healthcare risk waste / general waste?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "safe_system_manual_handling",
            "question_text": "4. Is there a safe system of work in place to minimise the manual handling risks associated with the segregation, disposal, and transportation of waste?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "clinical_waste_secured",
            "question_text": "5. Is clinical risk waste stored away from the public in a secured area with doors to secure area displaying a biohazard symbol and the wording 'no authorised entry/restricted access'?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "waste_bins_clean",
            "question_text": "6. Are all waste bins visibly clean, in good repair and included in a documented cleaning schedule?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "clinical_waste_regular_collection",
            "question_text": "7. Is clinical waste collected regularly to avoid build-up of waste?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "containment_system",
            "question_text": "8. Is a containment system such as a trolley or UN approved wheeled bin used to transport the clinical risk waste to the waste compound for offsite disposal?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "waste_training",
            "question_text": "9. Have staff who generate, segregate and package healthcare risk waste received appropriate training and training records maintained?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "ppe_provided",
            "question_text": "10. Is PPE provided based on Risk Assessment?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        # Use of Clinical Waste Bags
        {
            "id": "yellow_bags_marked",
            "question_text": "11. Are all large yellow bags marked UN 5H4 and display a class 6.2 label, the text UN 3291 and display the biohazard symbol and the words Clinical Waste?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "yellow_bags_placement",
            "question_text": "12. Are yellow bags placed in enclosed, pedal operated, lidded, non-combustible waste bins to minimise the risk of injury? (Please note they must not be tied onto containers/trolleys)",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "waste_not_decanted",
            "question_text": "13. Clinical risk waste is not decanted?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "bags_tied_properly",
            "question_text": "14. Are yellow bags tied appropriately with a swan necktie when ⅔ full?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "bin_holder_design",
            "question_text": "15. Is the bin holder front opening to facilitate ease of removal of a filled yellow bag and is it constructed in a way that facilitates effective cleaning?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "bin_holder_labeling",
            "question_text": "16. Does the bin holder list the permitted contents, display the biohazard symbol and text 'clinical risk waste'?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        # Clinical Risk Waste Collected in Rigid Bins
        {
            "id": "rigid_bins_assembly",
            "question_text": "17. Does the person assembling and closing the rigid bins comply with manufacturer's instructions on use?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "instructions_posted",
            "question_text": "18. Are posters with appropriate instructions on use located at bin assembly locations?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "lid_closed",
            "question_text": "19. Is the lid of the rigid bin closed when not in use?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "suitable_bin_size",
            "question_text": "20. Are rigid bins of a suitable size used to minimise length of time of use?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "bins_filled_properly",
            "question_text": "21. Are bins filled in accordance with manufacturers' guidelines i.e., bins are not filled beyond a maximum ⅔ full or at manufacturers fill line?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "bins_stable",
            "question_text": "22. Are bins stable and secured to prevent them from being inadvertently knocked over? e.g., Rigid bin holder",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        # Sharps
        {
            "id": "sharps_training",
            "question_text": "23. All staff have received the appropriate information, instruction, training, and supervision in the safe handling, use and disposal of sharps?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "safer_sharps",
            "question_text": "24. Is there any program to use of safer sharps – where available and when clinically practical?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "no_reheating_needles",
            "question_text": "25. No reheating of needles?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "sharps_containers_assembly",
            "question_text": "26. All sharps' containers are assembled correctly?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "enough_sharp_bins",
            "question_text": "27. Is there enough sharp bin at the point of use?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        # Incident Reporting
        {
            "id": "incident_reporting",
            "question_text": "28. Are all incidents where clinical risk waste is incorrectly presented for internal collection, reported in line with the local incident reporting procedures?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "incidents_recorded",
            "question_text": "29. Are all incidents/accidents/near misses recorded and investigated and remedial measures implemented?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        # Risk Assessments
        {
            "id": "risk_assessments_communicated",
            "question_text": "30. Are the results of the risk assessment communicated to all relevant employees and all who come into contact with HSE services and activities?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        {
            "id": "risk_assessments_reviewed",
            "question_text": "31. Are risk assessments reviewed at least annually or more frequently, if necessary, i.e., accident/incident or a change in circumstances to which they relate?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        },
        # Emergency Preparedness
        {
            "id": "emergency_training",
            "question_text": "32. Relevant staff have been trained and aware of how to act in the event of emergency?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No", "NA"]
        }
    ]


def is_waste_risk_assessment(file_path: str) -> bool:
    """
    Determine if the file is likely a Waste Risk Assessment Checklist.
    
    Args:
        file_path: The path to the file
        
    Returns:
        bool: True if the file appears to be a waste risk assessment
    """
    try:
        # If it's a file path, read content
        if isinstance(file_path, str) and os.path.exists(file_path):
            try:
                import docx
                doc = docx.Document(file_path)
                content = ""
                for para in doc.paragraphs:
                    content += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                content += paragraph.text + "\n"
                file_path = content
            except:
                # If we can't read as a docx, try to read as text
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        file_path = f.read()
                except:
                    pass
        
        # Check for key phrases that indicate it's a waste risk assessment
        lower_content = file_path.lower()
        
        # Define key markers that strongly indicate this form
        key_markers = [
            "waste risk assessment",
            "clinical risk waste",
            "healthcare risk waste",
            "sharps",
            "general waste",
            "waste segregation",
            "yellow bag",
            "biohazard symbol"
        ]
        
        # Count how many key markers are found
        marker_count = sum(1 for marker in key_markers if marker in lower_content)
        
        # If at least 3 key markers are found, it's likely this form
        return marker_count >= 3
        
    except Exception as e:
        print(f"Error checking if document is a waste risk assessment: {str(e)}")
        return False