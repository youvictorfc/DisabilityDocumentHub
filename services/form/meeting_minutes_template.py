"""
Template-based fallback solution for the Meeting Minutes form
This module provides a predefined field structure for the specific form
to ensure consistent extraction even when AI-based extraction fails.
"""

from typing import List, Dict, Any
import os


def get_meeting_minutes_template() -> List[Dict[str, Any]]:
    """
    Returns a predefined template for the Meeting Minutes form
    with the exact structure matching the document.
    """
    return [
        {
            "id": "meeting_subject",
            "question_text": "Subject OF MEETING/ TRAINING:",
            "field_type": "text",
            "required": True
        },
        {
            "id": "meeting_date",
            "question_text": "Date:",
            "field_type": "date",
            "required": True
        },
        {
            "id": "meeting_location",
            "question_text": "Held at:",
            "field_type": "text",
            "required": True
        },
        {
            "id": "meeting_time",
            "question_text": "Time:",
            "field_type": "text",
            "required": True
        },
        {
            "id": "supervisor_signature",
            "question_text": "NAME & SIGNATURE of supervisor or presenter:",
            "field_type": "text",
            "required": True
        },
        {
            "id": "topics_covered",
            "question_text": "TOPICS AND/OR ISSUES COVERED:",
            "field_type": "textarea",
            "required": True
        },
        # Attendees section - using a repeating group approach
        {
            "id": "attendee_name_1",
            "question_text": "NAME IN FULL (Attendee 1)",
            "field_type": "text",
            "required": False
        },
        {
            "id": "attendee_signature_1",
            "question_text": "SIGNATURE (Attendee 1)",
            "field_type": "signature",
            "required": False
        },
        {
            "id": "attendee_name_2",
            "question_text": "NAME IN FULL (Attendee 2)",
            "field_type": "text",
            "required": False
        },
        {
            "id": "attendee_signature_2",
            "question_text": "SIGNATURE (Attendee 2)",
            "field_type": "signature",
            "required": False
        },
        {
            "id": "attendee_name_3",
            "question_text": "NAME IN FULL (Attendee 3)",
            "field_type": "text",
            "required": False
        },
        {
            "id": "attendee_signature_3",
            "question_text": "SIGNATURE (Attendee 3)",
            "field_type": "signature",
            "required": False
        },
        {
            "id": "attendee_name_4",
            "question_text": "NAME IN FULL (Attendee 4)",
            "field_type": "text",
            "required": False
        },
        {
            "id": "attendee_signature_4",
            "question_text": "SIGNATURE (Attendee 4)",
            "field_type": "signature",
            "required": False
        },
        {
            "id": "attendee_name_5",
            "question_text": "NAME IN FULL (Attendee 5)",
            "field_type": "text",
            "required": False
        },
        {
            "id": "attendee_signature_5",
            "question_text": "SIGNATURE (Attendee 5)",
            "field_type": "signature",
            "required": False
        },
        {
            "id": "attendee_name_6",
            "question_text": "NAME IN FULL (Attendee 6)",
            "field_type": "text",
            "required": False
        },
        {
            "id": "attendee_signature_6",
            "question_text": "SIGNATURE (Attendee 6)",
            "field_type": "signature",
            "required": False
        },
        {
            "id": "attendee_name_7",
            "question_text": "NAME IN FULL (Attendee 7)",
            "field_type": "text",
            "required": False
        },
        {
            "id": "attendee_signature_7",
            "question_text": "SIGNATURE (Attendee 7)",
            "field_type": "signature",
            "required": False
        },
        {
            "id": "attendee_name_8",
            "question_text": "NAME IN FULL (Attendee 8)",
            "field_type": "text",
            "required": False
        },
        {
            "id": "attendee_signature_8",
            "question_text": "SIGNATURE (Attendee 8)",
            "field_type": "signature",
            "required": False
        }
    ]


def is_meeting_minutes(file_path: str) -> bool:
    """
    Determine if the file is likely a Meeting Minutes document.
    
    Args:
        file_path: The path to the file or its content
        
    Returns:
        bool: True if the file appears to be a meeting minutes document
    """
    try:
        # If it's a file path, read content
        if isinstance(file_path, str) and os.path.exists(file_path):
            try:
                import docx
                doc = docx.Document(file_path)
                content = ""
                for para in doc.paragraphs:
                    content += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                content += paragraph.text + "\n"
                file_path = content
            except:
                # If we can't read as a docx, try to read as text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    file_path = f.read()
        
        # Check for key phrases that indicate it's a meeting minutes document
        lower_content = file_path.lower()
        
        # Define key markers that strongly indicate this form
        key_markers = [
            "subject of meeting",
            "meeting/ training",
            "held at",
            "topics and/or issues covered",
            "name & signature of supervisor"
        ]
        
        # Count how many key markers are found
        marker_count = sum(1 for marker in key_markers if marker.lower() in lower_content)
        
        # If at least 2 key markers are found, it's likely this form
        return marker_count >= 2
        
    except Exception as e:
        print(f"Error checking if document is a meeting minutes: {str(e)}")
        return False