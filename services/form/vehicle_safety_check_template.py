"""
Template-based fallback solution for the Vehicle Safety Check Sheet
This module provides a predefined field structure for the specific form
to ensure consistent extraction even when AI-based extraction fails.
"""

from typing import List, Dict, Any
import os
import re


def get_vehicle_safety_check_template() -> List[Dict[str, Any]]:
    """
    Returns a predefined template for the Vehicle Safety Check Sheet
    with the exact structure matching the document.
    """
    return [
        {
            "id": "vehicle_registration",
            "question_text": "Vehicle registration #.:",
            "field_type": "text",
            "required": True,
            "options": []
        },
        {
            "id": "trip_date",
            "question_text": "Trip date:",
            "field_type": "date",
            "required": True,
            "options": []
        },
        {
            "id": "start_km",
            "question_text": "Start KM:",
            "field_type": "text",
            "required": True,
            "options": []
        },
        {
            "id": "end_km",
            "question_text": "End KM:",
            "field_type": "text",
            "required": True,
            "options": []
        },
        {
            "id": "drivers_name",
            "question_text": "Driver's name:",
            "field_type": "text",
            "required": True,
            "options": []
        },
        {
            "id": "drivers_licence",
            "question_text": "Driver's licence #",
            "field_type": "text",
            "required": True,
            "options": []
        },
        {
            "id": "driven_before",
            "question_text": "Have you ever driven this vehicle before?",
            "field_type": "radio",
            "required": True,
            "options": ["Yes", "No"]
        },
        {
            "id": "familiarise",
            "question_text": "If no, please familiarise yourself with the operator's manual and safety features of the vehicle.",
            "field_type": "textarea",
            "required": False,
            "options": []
        },
        # Safety Checklist Items - Days of the week as checkboxes
        # Lighting
        {
            "id": "lighting_mon",
            "question_text": "Lighting - Monday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "lighting_tue",
            "question_text": "Lighting - Tuesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "lighting_wed",
            "question_text": "Lighting - Wednesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "lighting_thu",
            "question_text": "Lighting - Thursday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "lighting_fri",
            "question_text": "Lighting - Friday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "lighting_sat",
            "question_text": "Lighting - Saturday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "lighting_sun",
            "question_text": "Lighting - Sunday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        # Vision
        {
            "id": "vision_mon",
            "question_text": "Vision - Monday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "vision_tue",
            "question_text": "Vision - Tuesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "vision_wed",
            "question_text": "Vision - Wednesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "vision_thu",
            "question_text": "Vision - Thursday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "vision_fri",
            "question_text": "Vision - Friday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "vision_sat",
            "question_text": "Vision - Saturday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "vision_sun",
            "question_text": "Vision - Sunday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        # Horn
        {
            "id": "horn_mon",
            "question_text": "Horn - Monday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "horn_tue",
            "question_text": "Horn - Tuesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "horn_wed",
            "question_text": "Horn - Wednesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "horn_thu",
            "question_text": "Horn - Thursday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "horn_fri",
            "question_text": "Horn - Friday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "horn_sat",
            "question_text": "Horn - Saturday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "horn_sun",
            "question_text": "Horn - Sunday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        # Brakes
        {
            "id": "brakes_mon",
            "question_text": "Brakes - Monday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "brakes_tue",
            "question_text": "Brakes - Tuesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "brakes_wed",
            "question_text": "Brakes - Wednesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "brakes_thu",
            "question_text": "Brakes - Thursday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "brakes_fri",
            "question_text": "Brakes - Friday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "brakes_sat",
            "question_text": "Brakes - Saturday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "brakes_sun",
            "question_text": "Brakes - Sunday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        # Wheel Assembly
        {
            "id": "wheels_mon",
            "question_text": "Wheel Assembly - Monday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "wheels_tue",
            "question_text": "Wheel Assembly - Tuesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "wheels_wed",
            "question_text": "Wheel Assembly - Wednesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "wheels_thu",
            "question_text": "Wheel Assembly - Thursday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "wheels_fri",
            "question_text": "Wheel Assembly - Friday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "wheels_sat",
            "question_text": "Wheel Assembly - Saturday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "wheels_sun",
            "question_text": "Wheel Assembly - Sunday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        # Fluid Levels
        {
            "id": "fluids_mon",
            "question_text": "Fluid Levels - Monday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "fluids_tue",
            "question_text": "Fluid Levels - Tuesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "fluids_wed",
            "question_text": "Fluid Levels - Wednesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "fluids_thu",
            "question_text": "Fluid Levels - Thursday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "fluids_fri",
            "question_text": "Fluid Levels - Friday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "fluids_sat",
            "question_text": "Fluid Levels - Saturday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "fluids_sun",
            "question_text": "Fluid Levels - Sunday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        # Visible Leaks
        {
            "id": "leaks_mon",
            "question_text": "Visible Leaks - Monday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "leaks_tue",
            "question_text": "Visible Leaks - Tuesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "leaks_wed",
            "question_text": "Visible Leaks - Wednesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "leaks_thu",
            "question_text": "Visible Leaks - Thursday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "leaks_fri",
            "question_text": "Visible Leaks - Friday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "leaks_sat",
            "question_text": "Visible Leaks - Saturday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "leaks_sun",
            "question_text": "Visible Leaks - Sunday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        # General Safety
        {
            "id": "safety_mon",
            "question_text": "General Safety - Monday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "safety_tue",
            "question_text": "General Safety - Tuesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "safety_wed",
            "question_text": "General Safety - Wednesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "safety_thu",
            "question_text": "General Safety - Thursday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "safety_fri",
            "question_text": "General Safety - Friday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "safety_sat",
            "question_text": "General Safety - Saturday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "safety_sun",
            "question_text": "General Safety - Sunday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        # General Cleanliness
        {
            "id": "cleanliness_mon",
            "question_text": "General Cleanliness - Monday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "cleanliness_tue",
            "question_text": "General Cleanliness - Tuesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "cleanliness_wed",
            "question_text": "General Cleanliness - Wednesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "cleanliness_thu",
            "question_text": "General Cleanliness - Thursday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "cleanliness_fri",
            "question_text": "General Cleanliness - Friday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "cleanliness_sat",
            "question_text": "General Cleanliness - Saturday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "cleanliness_sun",
            "question_text": "General Cleanliness - Sunday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        # Operating Check
        {
            "id": "operating_mon",
            "question_text": "Operating Check - Monday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "operating_tue",
            "question_text": "Operating Check - Tuesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "operating_wed",
            "question_text": "Operating Check - Wednesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "operating_thu",
            "question_text": "Operating Check - Thursday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "operating_fri",
            "question_text": "Operating Check - Friday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "operating_sat",
            "question_text": "Operating Check - Saturday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "operating_sun",
            "question_text": "Operating Check - Sunday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        # Fire Equipment
        {
            "id": "fire_mon",
            "question_text": "Fire Equipment - Monday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "fire_tue",
            "question_text": "Fire Equipment - Tuesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "fire_wed",
            "question_text": "Fire Equipment - Wednesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "fire_thu",
            "question_text": "Fire Equipment - Thursday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "fire_fri",
            "question_text": "Fire Equipment - Friday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "fire_sat",
            "question_text": "Fire Equipment - Saturday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "fire_sun",
            "question_text": "Fire Equipment - Sunday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        # Safety Equipment
        {
            "id": "safety_equipment_mon",
            "question_text": "Safety Equipment - Monday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "safety_equipment_tue",
            "question_text": "Safety Equipment - Tuesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "safety_equipment_wed",
            "question_text": "Safety Equipment - Wednesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "safety_equipment_thu",
            "question_text": "Safety Equipment - Thursday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "safety_equipment_fri",
            "question_text": "Safety Equipment - Friday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "safety_equipment_sat",
            "question_text": "Safety Equipment - Saturday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "safety_equipment_sun",
            "question_text": "Safety Equipment - Sunday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        # Communications
        {
            "id": "comms_mon",
            "question_text": "Communications - Monday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "comms_tue",
            "question_text": "Communications - Tuesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "comms_wed",
            "question_text": "Communications - Wednesday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "comms_thu",
            "question_text": "Communications - Thursday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "comms_fri",
            "question_text": "Communications - Friday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "comms_sat",
            "question_text": "Communications - Saturday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "comms_sun",
            "question_text": "Communications - Sunday",
            "field_type": "checkbox",
            "required": False,
            "options": ["Satisfactory", "Defect"]
        },
        {
            "id": "defect_info",
            "question_text": "Please provide information about the defect:",
            "field_type": "textarea",
            "required": False,
            "options": []
        },
        {
            "id": "driver_declaration",
            "question_text": "Driver's declaration: I have inspected the vehicle as required and to the best of my knowledge the vehicle is in a suitable equipped and safe condition. I declare myself in a fit state to drive this vehicle. Driver to acknowledge above statement by completing the below:",
            "field_type": "text",
            "required": True,
            "options": []
        },
        {
            "id": "driver_name",
            "question_text": "Driver's name:",
            "field_type": "text",
            "required": True,
            "options": []
        },
        {
            "id": "driver_signature_date",
            "question_text": "Date:",
            "field_type": "date",
            "required": True,
            "options": []
        }
    ]


def is_vehicle_safety_check(file_path: str) -> bool:
    """
    Determine if the file is likely a Vehicle Safety Check Sheet.
    
    Args:
        file_path: The path to the file
        
    Returns:
        bool: True if the file appears to be a vehicle safety check
    """
    try:
        # If it's a file path, read content
        if isinstance(file_path, str) and os.path.exists(file_path):
            try:
                import docx
                doc = docx.Document(file_path)
                content = ""
                for para in doc.paragraphs:
                    content += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                content += paragraph.text + "\n"
                file_path = content
            except:
                # If we can't read as a docx, try to read as text
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        file_path = f.read()
                except:
                    pass
        
        # Check for key phrases that indicate it's a vehicle safety check
        lower_content = file_path.lower()
        
        # Define key markers that strongly indicate this form
        key_markers = [
            "vehicle safety check",
            "have you ever driven this vehicle before",
            "driver's declaration",
            "lighting",
            "vision",
            "horn",
            "brakes",
            "wheel assembly",
            "fluid levels",
            "visible leaks"
        ]
        
        # Count how many key markers are found
        marker_count = sum(1 for marker in key_markers if marker in lower_content)
        
        # If at least 3 key markers are found, it's likely this form
        return marker_count >= 3
        
    except Exception as e:
        print(f"Error checking if document is a vehicle safety check: {str(e)}")
        return False